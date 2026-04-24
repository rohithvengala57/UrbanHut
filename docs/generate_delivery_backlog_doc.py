from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "delivery-backlog.md"
OUTPUT = ROOT / "urban-hut-delivery-backlog.docx"


def add_bullet(document: Document, text: str, level: int = 0) -> None:
    style = "List Bullet" if level == 0 else "List Bullet 2"
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)


def add_number(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(text)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(text)


def add_heading(document: Document, text: str, level: int) -> None:
    document.add_heading(text, level=level)


def format_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 3"].font.name = "Aptos Display"
    styles["Heading 4"].font.name = "Aptos Display"


def render_markdown(document: Document, markdown: str) -> None:
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()

        if not line:
            continue

        if line.startswith("# "):
            add_heading(document, line[2:].strip(), 0)
            continue
        if line.startswith("## "):
            add_heading(document, line[3:].strip(), 1)
            continue
        if line.startswith("### "):
            add_heading(document, line[4:].strip(), 2)
            continue
        if line.startswith("#### "):
            add_heading(document, line[5:].strip(), 3)
            continue
        if line.startswith("- `") or line.startswith("- "):
            add_bullet(document, line[2:].strip(), 0)
            continue
        if line.startswith("1. "):
            add_number(document, line[3:].strip())
            continue
        if line.startswith("---"):
            document.add_paragraph("_" * 60)
            continue

        add_body_paragraph(document, line)


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    document = Document()
    format_document(document)
    render_markdown(document, markdown)

    title = document.paragraphs[0]
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if title.runs:
        title.runs[0].font.size = Pt(20)
        title.runs[0].bold = True

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
